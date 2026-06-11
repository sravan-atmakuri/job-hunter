"""
agent_resume.py — Reads jobs_filtered.json and the original resume, tailors
resume keywords to each job description using the Claude CLI, saves each
tailored resume as output/resumes/resume_{company}_{id}.docx, and logs the
mapping to output/resume_map.json.

Strategy: copy the original .docx and do targeted in-place edits so all
fonts, layout, and formatting are perfectly preserved.

Uses claude -p via claude_cli.py — no API key needed.
"""

from __future__ import annotations

import json
import logging
import re
import shutil
import time
from pathlib import Path

import yaml
from docx import Document

from claude_cli import call_claude_json

logging.basicConfig(level=logging.INFO, format="%(asctime)s [resume] %(message)s")
log = logging.getLogger(__name__)

INPUT_FILE = Path("output/jobs_filtered.json")
RESUME_MAP_FILE = Path("output/resume_map.json")
RESUME_INDEX_FILE = Path("output/resume_index.json")
RESUMES_DIR = Path("output/resumes")
CONFIG_FILE = Path("config.yaml")


def load_config() -> dict:
    with open(CONFIG_FILE) as f:
        return yaml.safe_load(f)


# ---------------------------------------------------------------------------
# Select the best base resume for a given job title
# ---------------------------------------------------------------------------

def select_resume(job: dict, config: dict) -> str:
    """
    Pick the right base resume from config resume_map based on job title keywords.
    Falls back to config resume_path if no match.
    """
    resume_map: dict = config.get("resume_map", {})
    title_lower = job.get("title", "").lower()

    # Salesforce developer/admin roles → salesforce resume
    if any(kw in title_lower for kw in ["salesforce developer", "sfdc developer", "salesforce admin"]):
        if "salesforce" in resume_map:
            path = resume_map["salesforce"]
            if Path(path).exists():
                return path

    # DataCloud roles → datacloud resume
    if any(kw in title_lower for kw in ["datacloud", "data cloud"]):
        if "datacloud" in resume_map:
            path = resume_map["datacloud"]
            if Path(path).exists():
                return path

    # Salesforce QA / test roles that mention Salesforce → qa resume (has both)
    if "salesforce" in title_lower:
        if "qa" in resume_map:
            path = resume_map["qa"]
            if Path(path).exists():
                return path

    # QA / Test / Quality roles → qa resume
    if any(kw in title_lower for kw in ["qa", "qe", "quality", "test"]):
        if "qa" in resume_map:
            path = resume_map["qa"]
            if Path(path).exists():
                return path

    # Ultimate fallback
    return config.get("resume_path", "")


# ---------------------------------------------------------------------------
# Resume reader (for passing text to Claude)
# ---------------------------------------------------------------------------

def read_resume(resume_path: str) -> str:
    path = Path(resume_path)
    if not path.exists():
        log.warning("Resume file not found: %s", resume_path)
        return ""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ---------------------------------------------------------------------------
# Claude-based resume tailor
# ---------------------------------------------------------------------------

ANALYZE_AND_TAILOR_PROMPT = """ATS resume optimizer. Single-pass: analyze gaps then tailor. Target 75%+ match.

Resume skeleton:
{resume_skeleton}

JD ({title} @ {company}):
{jd}

Return JSON only (no markdown):
{{
  "match_percentage": <0-100 current match>,
  "matched_keywords": ["up to 15 keywords already in resume"],
  "missing_keywords": ["ALL keywords in JD absent from resume"],
  "top_skills": ["top 5 skills the role requires"],
  "recommendation": "1-sentence ATS tip",
  "summary": "2-3 sentence summary weaving in ≥3 missing keywords",
  "job1_bullets": ["3-4 Key Responsibilities bullets for most recent job, each with ≥1 missing keyword"],
  "job1_tech_bullets": ["2-3 Technical Contributions bullets for most recent job, aligning with job1_bullets"],
  "job2_bullets": ["2-3 Key Responsibilities bullets for second job, each with ≥1 missing keyword"],
  "job2_tech_bullets": ["1-2 Technical Contributions bullets for second job, aligning with job2_bullets"],
  "skills_to_add": ["every missing keyword not covered by the bullets/summary above"]
}}
Rules: facts only, bullets start with action verb, no "I", cover ALL missing keywords across all fields."""




def tailor_resume(resume_text: str, job: dict, missing_keywords: list | None = None) -> dict:
    """Single Claude call: ATS analysis + tailoring combined."""
    jd = job.get("full_description") or job.get("description_snippet", "")

    prompt = ANALYZE_AND_TAILOR_PROMPT.format(
        resume_skeleton=resume_text[:4000],  # full resume content for accurate keyword matching
        title=job["title"],
        company=job["company"],
        jd=jd[:6000],
    )
    result = call_claude_json(prompt)
    return result if isinstance(result, dict) else {}


# ---------------------------------------------------------------------------
# DOCX in-place modifier — copies original and patches specific paragraphs
# ---------------------------------------------------------------------------

# Markers that signal the start of the Professional Summary section
SUMMARY_MARKERS = {
    "professional summary", "summary", "objective", "professional profile", "profile",
}

# Markers that signal the start of a Skills / Core Competencies section
SKILLS_MARKERS = {
    "skills", "core competencies", "technical skills", "key skills",
    "core skills", "competencies", "areas of expertise",
}


def _clear_paragraph_text(para) -> None:
    """Remove all text from a paragraph while keeping its style/formatting."""
    for run in para.runs:
        run.text = ""


def _set_paragraph_text(para, new_text: str) -> None:
    """
    Replace the paragraph's text content, preserving the first run's formatting
    (font, size, bold, color, etc.).
    """
    if not para.runs:
        para.add_run(new_text)
        return

    # Put all new text in the first run, clear the rest
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _para_is_section_heading(para, markers: set) -> bool:
    """Return True if this paragraph's text matches one of the given section markers."""
    text = para.text.strip().lower()
    # Exact match or starts-with (handles "SKILLS & CERTIFICATIONS" etc.)
    return any(text == m or text.startswith(m) for m in markers)


def _insert_paragraph_after(ref_para, text: str, style=None) -> None:
    """Insert a new paragraph with `text` immediately after `ref_para`, copying its style."""
    from docx.oxml.ns import qn
    import copy
    new_p = copy.deepcopy(ref_para._p)
    # Clear runs and set new text in the first run
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    # Build a minimal run
    r_elem = copy.deepcopy(ref_para._p.findall(qn("w:r"))[0]) if ref_para._p.findall(qn("w:r")) else None
    if r_elem is not None:
        for t in r_elem.findall(qn("w:t")):
            r_elem.remove(t)
        from lxml import etree
        t_elem = etree.SubElement(r_elem, qn("w:t"))
        t_elem.text = text
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        new_p.append(r_elem)
    ref_para._p.addnext(new_p)


def _find_tech_contribution_insert_point(paras: list, job_num: int):
    """
    Find the last bullet paragraph in job N's Technical Contributions block.
    Returns (paragraph, index) or (None, None).
    """
    TECH_STARTS = ("technical contribution", "technical contributions", "technical achievements")
    EXP_SECTION = {"professional experience", "experience", "work experience"}
    STOP_SECTION = {"education", "certifications", "technical projects", "projects", "skills", "awards"}
    NEXT_RESP = ("key responsibilities", "responsibilities & achievements", "responsibilities:")

    tech_indices = []
    in_exp = False
    for i, para in enumerate(paras):
        lower = para.text.strip().lower()
        if not in_exp:
            if any(m in lower for m in EXP_SECTION):
                in_exp = True
            continue
        if any(m in lower for m in STOP_SECTION) and para.text.strip().isupper():
            break
        if any(lower.startswith(t) for t in TECH_STARTS):
            tech_indices.append(i)

    if not tech_indices:
        for i, para in enumerate(paras):
            lower = para.text.strip().lower()
            if any(lower.startswith(t) for t in TECH_STARTS):
                tech_indices.append(i)

    if job_num > len(tech_indices):
        return None, None

    start_idx = tech_indices[job_num - 1] + 1
    next_tech_idx = tech_indices[job_num] if job_num < len(tech_indices) else len(paras)

    last_bullet = None
    last_bullet_idx = None
    for i in range(start_idx, min(next_tech_idx, len(paras))):
        text = paras[i].text.strip()
        lower = text.lower()
        if any(lower.startswith(r) for r in NEXT_RESP):
            break
        if any(m in lower for m in STOP_SECTION) and text.isupper():
            break
        if text:
            last_bullet = paras[i]
            last_bullet_idx = i

    return last_bullet, last_bullet_idx


def _find_job_bullet_insert_point(paras: list, job_num: int):
    """
    Find the last bullet paragraph in job N's first "Key Responsibilities" block.
    Strategy: find all "Key Responsibilities" headings; pick the Nth one (1-based),
    then collect bullets until a sub-section header or blank separator.
    Returns (paragraph, index) or (None, None).
    """
    RESP_STARTS = ("key responsibilities", "responsibilities & achievements", "responsibilities:")
    SUB_SECTION = ("technical contribution", "project highlight", "project name", "key achievements")
    EXP_SECTION = {"professional experience", "experience", "work experience"}
    STOP_SECTION = {"education", "certifications", "technical projects", "projects", "skills", "awards"}

    # Step 1: find all "Key Responsibilities" paragraph indices.
    # Use a two-pass approach: first try the strict in_exp guard (works for QA resume),
    # then fall back to a global scan if nothing found (handles resumes where the
    # experience header is embedded/concatenated into another paragraph).
    resp_indices = []
    in_exp = False
    for i, para in enumerate(paras):
        lower = para.text.strip().lower()
        if not in_exp:
            # Match if EXP_SECTION keyword appears anywhere in the paragraph text
            if any(m in lower for m in EXP_SECTION):
                in_exp = True
            continue
        if any(m in lower for m in STOP_SECTION) and para.text.strip().isupper():
            break
        if any(lower.startswith(r) for r in RESP_STARTS):
            resp_indices.append(i)

    # Fallback 1: if experience section header was never found (e.g. embedded in another line),
    # scan all paragraphs globally for "Key Responsibilities" headings.
    if not resp_indices:
        for i, para in enumerate(paras):
            lower = para.text.strip().lower()
            if any(m in lower for m in STOP_SECTION) and para.text.strip().isupper():
                break
            if any(lower.startswith(r) for r in RESP_STARTS):
                resp_indices.append(i)

    # Fallback 2: for project-based resumes (e.g. DataCloud) that use "Project:" as
    # job block separators instead of "Key Responsibilities" headings.
    # In this mode, also stop bullet collection at blank-line job boundaries.
    stop_at_blank = False
    if not resp_indices:
        PROJ_STARTS = ("project:", "projects:")
        stop_at_blank = True
        for i, para in enumerate(paras):
            lower = para.text.strip().lower()
            if any(m in lower for m in STOP_SECTION) and para.text.strip().isupper():
                break
            if any(lower.startswith(r) for r in PROJ_STARTS):
                resp_indices.append(i)

    if job_num > len(resp_indices):
        return None, None

    # Step 2: starting after the target heading, collect bullets until a
    # sub-section break, next heading, or (in project mode) a blank line separator.
    start_idx = resp_indices[job_num - 1] + 1
    next_resp_idx = resp_indices[job_num] if job_num < len(resp_indices) else len(paras)

    last_bullet = None
    last_bullet_idx = None

    for i in range(start_idx, min(next_resp_idx, len(paras))):
        text = paras[i].text.strip()
        lower = text.lower()
        if any(lower.startswith(s) for s in SUB_SECTION):
            break
        if stop_at_blank and not text:
            break  # blank line = job boundary in project-based resumes
        if text:
            last_bullet = paras[i]
            last_bullet_idx = i

    return last_bullet, last_bullet_idx


def patch_docx(base_path: str, tailored: dict, job: dict, output_path: Path) -> None:
    """
    Copy base_path to output_path, then open the copy and:
    - Replace Professional Summary text
    - Append missing ATS skills to the Skills / Core Competencies section
    - Append ATS bullet(s) to job 1 and job 2 Key Responsibilities
    Preserves all original formatting.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(base_path, str(output_path))

    doc = Document(str(output_path))
    paras = doc.paragraphs

    summary_text = tailored.get("summary", "").strip()
    job1_bullets: list = tailored.get("job1_bullets", [])
    job2_bullets: list = tailored.get("job2_bullets", [])
    job1_tech_bullets: list = tailored.get("job1_tech_bullets", [])
    job2_tech_bullets: list = tailored.get("job2_tech_bullets", [])
    skills_to_add: list = [s.strip() for s in tailored.get("skills_to_add", []) if s.strip()]

    in_summary = False
    summary_done = False
    skills_done = False

    for para in paras:
        text_stripped = para.text.strip()

        if _para_is_section_heading(para, SUMMARY_MARKERS):
            in_summary = True
            continue

        if _para_is_section_heading(para, SKILLS_MARKERS):
            in_summary = False
            # Next non-empty paragraph is the skills content — patch it
            if skills_to_add and not skills_done:
                # We'll handle it on the NEXT paragraph
                in_summary = False
                # Use a flag to catch the next content paragraph
                _patch_skills_next = [True]  # mutable flag
                continue
            continue

        if in_summary and not summary_done and text_stripped and summary_text:
            _set_paragraph_text(para, summary_text)
            summary_done = True
            in_summary = False
            continue

    # Skills patching: find first content paragraph under skills heading and append
    if skills_to_add and not skills_done:
        in_skills = False
        for para in doc.paragraphs:
            text_stripped = para.text.strip()
            if _para_is_section_heading(para, SKILLS_MARKERS):
                in_skills = True
                continue
            if in_skills and text_stripped:
                existing = para.text.strip()
                # Only add skills not already in the text (case-insensitive)
                existing_lower = existing.lower()
                new_skills = [s for s in skills_to_add if s.lower() not in existing_lower]
                if new_skills:
                    separator = " | " if " | " in existing or "|" in existing else ", "
                    updated = existing + separator + separator.join(new_skills)
                    _set_paragraph_text(para, updated)
                    log.info("  Skills section updated: +%d keyword(s)", len(new_skills))
                skills_done = True
                break

    # Inject ATS bullets into Key Responsibilities for job 1 and job 2
    paras = doc.paragraphs

    for job_num, new_bullets in ((1, job1_bullets), (2, job2_bullets)):
        if not new_bullets:
            continue
        ref_para, ref_idx = _find_job_bullet_insert_point(paras, job_num)
        if ref_para is None:
            log.warning("  Could not locate job %d Key Responsibilities insert point", job_num)
            continue
        for bullet in reversed(new_bullets):
            bullet = bullet.lstrip("-•· ").strip()
            if bullet:
                _insert_paragraph_after(ref_para, bullet)
        log.debug("  Injected %d responsibility bullet(s) into job %d", len(new_bullets), job_num)
        paras = doc.paragraphs

    # Inject Technical Contributions bullets for job 1 and job 2
    for job_num, tech_bullets in ((1, job1_tech_bullets), (2, job2_tech_bullets)):
        if not tech_bullets:
            continue
        ref_para, ref_idx = _find_tech_contribution_insert_point(paras, job_num)
        if ref_para is None:
            log.warning("  Could not locate job %d Technical Contributions insert point", job_num)
            continue
        for bullet in reversed(tech_bullets):
            bullet = bullet.lstrip("-•· ").strip()
            if bullet:
                _insert_paragraph_after(ref_para, bullet)
        log.debug("  Injected %d tech contribution bullet(s) into job %d", len(tech_bullets), job_num)
        paras = doc.paragraphs

    doc.save(str(output_path))
    log.info("  → Saved: %s", output_path.name)

    if not summary_done:
        log.warning("  Could not patch summary in %s", Path(base_path).name)




# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def run(jobs: list[dict] | None = None) -> dict:
    config = load_config()

    if jobs is None:
        if not INPUT_FILE.exists():
            log.error("Input file not found: %s — run agent_relevance first", INPUT_FILE)
            return {}
        jobs = json.loads(INPUT_FILE.read_text())

    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    resume_map: dict[str, dict] = {}

    for i, job in enumerate(jobs, 1):
        log.info(
            "[%d/%d] Tailoring for '%s' @ %s",
            i, len(jobs), job["title"], job["company"],
        )

        base_resume = select_resume(job, config)
        if not base_resume or not Path(base_resume).exists():
            log.warning("  No valid base resume found for %s — skipping", job["title"])
            continue

        log.info("  Base resume: %s", Path(base_resume).name)

        # Keep the original resume filename; store in a per-job subdirectory: "Job Title - Company"
        original_filename = Path(base_resume).name
        raw_folder = f"{job['title']} - {job['company']}"
        folder_name = re.sub(r'[<>:"/\\|?*]', "_", raw_folder).strip()
        output_path = RESUMES_DIR / folder_name / original_filename

        if output_path.exists():
            log.info("  → Already exists, skipping")
        else:
            resume_text = read_resume(base_resume)
            tailored = tailor_resume(resume_text, job)
            if tailored:
                patch_docx(base_resume, tailored, job, output_path)
            else:
                log.warning("  Claude returned empty tailoring — copying base resume as-is")
                shutil.copy2(base_resume, str(output_path))
            time.sleep(1)

        resume_map[job["id"]] = {
            "job_id": job["id"],
            "job_title": job["title"],
            "company": job["company"],
            "job_url": job["url"],
            "resume_path": str(output_path),
            "relevance_score": job.get("relevance_score", 0),
            "keywords": job.get("keywords", []),
            "location": job["location"],
        }

    RESUME_MAP_FILE.parent.mkdir(parents=True, exist_ok=True)
    RESUME_MAP_FILE.write_text(json.dumps(resume_map, indent=2))
    log.info("Resume map saved: %d entries → %s", len(resume_map), RESUME_MAP_FILE)

    # Write human-readable index: job_id → resume filename + job info
    resume_index = {
        job_id: {
            "resume_filename": Path(entry["resume_path"]).name,
            "resume_path": entry["resume_path"],
            "job_title": entry["job_title"],
            "company": entry["company"],
            "job_url": entry["job_url"],
        }
        for job_id, entry in resume_map.items()
    }
    RESUME_INDEX_FILE.write_text(json.dumps(resume_index, indent=2))
    log.info("Resume index saved → %s", RESUME_INDEX_FILE)

    return resume_map


if __name__ == "__main__":
    run()
