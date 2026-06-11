"""
app.py — Flask web UI for ATS resume tailoring + job auto-apply.

Usage:
    python3 app.py
    then open http://localhost:5001
"""
from __future__ import annotations

import json
import logging
import re
import threading
import time
from pathlib import Path
from queue import Empty, Queue
from urllib.parse import urlparse

import yaml
from flask import Flask, Response, jsonify, render_template, request, send_file, stream_with_context

# Import existing agent functions
from agent_resume import SUMMARY_MARKERS, patch_docx, read_resume, tailor_resume
from agent_applier import PLATFORM_HANDLERS, apply_generic, detect_platform, get_applicant
from claude_cli import call_claude_json

logging.basicConfig(level=logging.INFO, format="%(asctime)s [app] %(message)s")
log = logging.getLogger(__name__)

CONFIG_FILE = Path("config.yaml")
UPLOAD_DIR = Path("output/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

app = Flask(__name__)

# Active apply sessions: session_id → Queue
_apply_sessions: dict[str, Queue] = {}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_config() -> dict:
    with open(CONFIG_FILE) as f:
        return yaml.safe_load(f)


def extract_original_summary(resume_path: str) -> str:
    """Pull the current Professional Summary paragraph from the docx."""
    try:
        from docx import Document
        doc = Document(resume_path)
        in_summary = False
        for para in doc.paragraphs:
            text = para.text.strip()
            lower = text.lower()
            if lower in SUMMARY_MARKERS or lower.startswith("professional summary"):
                in_summary = True
                continue
            if in_summary and text:
                return text
    except Exception:
        pass
    return ""


def extract_original_bullets(resume_path: str) -> tuple[list[str], list[str], list[str], list[str]]:
    """Return last 3 bullets from Key Responsibilities and Technical Contributions for job1 and job2."""
    try:
        from docx import Document
        doc = Document(resume_path)
        paras = doc.paragraphs

        RESP_STARTS = ("key responsibilities", "responsibilities & achievements", "responsibilities:")
        TECH_STARTS = ("technical contribution", "technical contributions", "technical achievements")
        EXP_MARKERS = {"professional experience", "experience", "work experience"}
        STOP_MARKERS = {"education", "certifications", "technical projects", "projects", "skills", "awards"}

        resp_indices: list[int] = []
        tech_indices: list[int] = []
        in_exp = False
        for i, para in enumerate(paras):
            lower = para.text.strip().lower()
            if not in_exp:
                if any(m in lower for m in EXP_MARKERS):
                    in_exp = True
                continue
            if any(m in lower for m in STOP_MARKERS) and para.text.strip().isupper():
                break
            if any(lower.startswith(r) for r in RESP_STARTS):
                resp_indices.append(i)
            elif any(lower.startswith(t) for t in TECH_STARTS):
                tech_indices.append(i)

        if not resp_indices:
            for i, para in enumerate(paras):
                lower = para.text.strip().lower()
                if any(lower.startswith(r) for r in RESP_STARTS):
                    resp_indices.append(i)
        if not tech_indices:
            for i, para in enumerate(paras):
                lower = para.text.strip().lower()
                if any(lower.startswith(t) for t in TECH_STARTS):
                    tech_indices.append(i)

        def _collect(indices, job_num, stop_starts):
            bucket: list[str] = []
            if job_num > len(indices):
                return bucket
            start = indices[job_num - 1] + 1
            end = indices[job_num] if job_num < len(indices) else len(paras)
            for i in range(start, min(end, len(paras))):
                text = paras[i].text.strip()
                lower = text.lower()
                if text and not any(lower.startswith(s) for s in stop_starts):
                    bucket.append(text)
                    if len(bucket) >= 4:
                        break
            return bucket

        resp_stop = ("technical", "project")
        tech_stop = ("key resp", "responsibilities")

        b1 = _collect(resp_indices, 1, resp_stop)[-3:]
        b2 = _collect(resp_indices, 2, resp_stop)[-3:]
        t1 = _collect(tech_indices, 1, tech_stop)[-3:]
        t2 = _collect(tech_indices, 2, tech_stop)[-3:]

        return b1, b2, t1, t2
    except Exception:
        return [], [], [], []


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/config")
def api_config():
    config = load_config()
    resume_map = config.get("resume_map", {})
    resumes = {}
    labels = {
        "qa": "QA / Test Engineer",
        "salesforce": "Salesforce Developer",
        "datacloud": "DataCloud Specialist",
    }
    for key, path in resume_map.items():
        p = Path(path)
        resumes[key] = {
            "path": str(p),
            "exists": p.exists(),
            "name": p.name,
            "label": labels.get(key, key.title()),
        }

    applicant = config.get("applicant", {})
    return jsonify({
        "resumes": resumes,
        "applicant": {
            "name": f"{applicant.get('first_name', '')} {applicant.get('last_name', '')}".strip(),
            "email": applicant.get("email", ""),
            "linkedin": applicant.get("linkedin_url", ""),
            "phone": applicant.get("phone", ""),
        },
        "default_resume": "qa",
    })


def _scrape_with_playwright(url: str, domain: str) -> dict:
    """Use headless Playwright to scrape JS-rendered job pages (LinkedIn, Dice, etc.)."""
    try:
        from playwright.sync_api import sync_playwright
        SESSION_FILE = Path("output/linkedin_session.json")

        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            ctx_kwargs: dict = {
                "user_agent": (
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
                ),
            }
            if "linkedin.com" in domain and SESSION_FILE.exists():
                ctx_kwargs["storage_state"] = str(SESSION_FILE)
                log.info("LinkedIn: using saved session for scrape")

            context = browser.new_context(**ctx_kwargs)
            page = context.new_page()
            page.goto(url, wait_until="domcontentloaded", timeout=30000)

            title = company = description = ""

            if "linkedin.com" in domain:
                # Give React time to hydrate the job page
                page.wait_for_timeout(5000)

                # Title + company from <title>: "Job Title | Company | LinkedIn"
                page_title = page.title()
                parts = [p.strip() for p in page_title.split("|")]
                if len(parts) >= 2:
                    title = parts[0]
                    company = parts[1] if parts[1].lower() != "linkedin" else ""

                # --- Strategy 1: target the job description container directly ---
                # LinkedIn renders the description inside one of these containers.
                # innerText on the container avoids all footer/sidebar/similar-jobs noise.
                description = ""
                for sel in [
                    ".jobs-description__content",
                    ".jobs-description-content__text",
                    ".jobs-box__html-content",
                    "div[class*='jobs-description__content']",
                    "div[class*='description__text']",
                    "#job-details",
                ]:
                    try:
                        el = page.query_selector(sel)
                        if el:
                            raw = el.inner_text().strip()
                            if len(raw) > 200:
                                description = re.sub(r"\n{3,}", "\n\n", raw)
                                log.info("LinkedIn description via CSS selector %r (%d chars)", sel, len(description))
                                break
                    except Exception:
                        pass

                # --- Strategy 2: text parsing — find "About the job" block ---
                if not description:
                    body_text = page.evaluate("() => document.body.innerText") or ""
                    body_text = re.sub(r"\n{3,}", "\n\n", body_text)

                    desc_start = -1
                    for marker in ["about the job\n", "job description\n", "about this role\n"]:
                        idx = body_text.lower().find(marker)
                        if idx != -1:
                            desc_start = idx + len(marker)
                            break

                    if desc_start != -1:
                        desc_raw = body_text[desc_start:]
                        # Trim at the first line that signals end of the job description block
                        end_markers = [
                            "\nShow less\n", "\nSee less\n",
                            "\nReport this job\n", "\nReport job\n",
                            "\nMeet the hiring team\n",
                            "\nAbout the company\n",
                            "\nSimilar jobs\n", "\nMore jobs like this\n",
                            "\nPeople also viewed\n",
                            "\nShow more jobs\n",
                            "\nGet job alerts\n",
                            "\nSet alert for similar jobs\n",
                            "\nSet alert\n",
                            "\nJob search faster with Premium\n",
                            "\nJob search faster\n",
                            "\nReactivate Premium\n",
                            "\nAccess company insights\n",
                            "\nuse Premium\n",
                            "\nCancel anytime\n",
                            "\nNo hidden fees\n",
                            "\nmembers use Premium\n",
                        ]
                        best_end = len(desc_raw)
                        for em in end_markers:
                            idx = desc_raw.lower().find(em.lower())
                            if 0 < idx < best_end:
                                best_end = idx
                        description = desc_raw[:best_end].strip()
                        log.info("LinkedIn description via text parse (%d chars)", len(description))

            elif "dice.com" in domain:
                try:
                    page.wait_for_selector('[data-cy="jobTitle"], h1[class*="title"]', timeout=12000)
                    page.wait_for_timeout(1500)
                except Exception:
                    pass

                for sel in ['[data-cy="jobTitle"]', 'h1[class*="title"]', "h1"]:
                    el = page.query_selector(sel)
                    if el:
                        title = el.inner_text().strip()
                        break

                for sel in ['[data-cy="companyNameLink"]', '[data-cy="companyName"]', 'a[class*="company"]']:
                    el = page.query_selector(sel)
                    if el:
                        company = el.inner_text().strip()
                        break

                for sel in ['[data-cy="jobDescription"]', ".job-description", "#jobdescSec"]:
                    el = page.query_selector(sel)
                    if el:
                        raw = el.inner_text().strip()
                        from bs4 import BeautifulSoup as _BS
                        description = _BS(raw, "html.parser").get_text(separator="\n").strip()
                        description = re.sub(r"\n{3,}", "\n\n", description)
                        break

            else:
                # Generic JS page — grab all visible body text
                page.wait_for_timeout(3000)
                body_text = page.evaluate("() => document.body.innerText")
                description = re.sub(r"\n{3,}", "\n\n", body_text.strip())[:6000]

            page.close()
            browser.close()

        # Detect feed/listing/not-found pages instead of a real job posting
        not_found_signals = ["page not found", "uh oh, we can", "strànka nenalezena", "seite nicht gefunden"]
        listing_titles = {"jobs", "linkedin jobs", "job search", "linkedin", ""}
        # Listing page signals: no title, generic title, or notification-count prefix like "(39) Top job picks"
        is_listing = (
            title.lower() in listing_titles
            or bool(re.match(r"^\(\d+\)", title))   # "(39) Top job picks for you"
            or "top job picks" in title.lower()
            or "job picks for you" in title.lower()
            or "sign in" in title.lower()
            or "log in" in title.lower()
        )
        if is_listing or (description and any(sig in description.lower()[:500] for sig in not_found_signals)):
            log.warning("Playwright: detected non-job page for %s (title=%r)", url, title)
            description = ""
            title = ""
            company = ""

        return {"title": title, "company": company, "description": description}

    except Exception as exc:
        log.warning("Playwright scrape failed for %s: %s", domain, exc)
        return {"title": "", "company": "", "description": "", "error": str(exc)}


def _scrape_with_requests(url: str) -> dict:
    """Fast BeautifulSoup scraper — works for server-rendered pages (Greenhouse, Lever, etc.)."""
    try:
        import requests as req_lib
        from bs4 import BeautifulSoup

        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
            ),
            "Accept-Language": "en-US,en;q=0.9",
        }
        resp = req_lib.get(url, headers=headers, timeout=20, allow_redirects=True)
        soup = BeautifulSoup(resp.text, "html.parser")

        # Fast path: JSON-LD structured data (many job boards include this)
        import json as _json
        for tag in soup.find_all("script", type="application/ld+json"):
            try:
                d = _json.loads(tag.string or "")
                # Could be a list or a single object
                items = d if isinstance(d, list) else [d]
                for item in items:
                    if item.get("@type") in ("JobPosting", "jobPosting") and item.get("title"):
                        raw_desc = item.get("description", "")
                        desc_text = BeautifulSoup(raw_desc, "html.parser").get_text(separator="\n") if "<" in raw_desc else raw_desc
                        return {
                            "title": item.get("title", ""),
                            "company": item.get("hiringOrganization", {}).get("name", ""),
                            "description": desc_text.strip(),
                        }
            except Exception:
                pass

        # Fall back: remove noise tags then extract visible text
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text)

        # If page looks like a login wall or "not found", bail
        lower_text = text.lower()
        if len(text) < 500 or ("page not found" in lower_text and len(text) < 2000):
            return {"title": "", "company": "", "description": ""}

        # Ask Claude to extract structured fields from the raw text
        prompt = (
            "Extract job details from this job page text. Return JSON only (no markdown):\n"
            '{"title":"exact job title","company":"company name","description":"full job description with requirements and responsibilities"}\n\n'
            f"Page URL: {url}\nPage text:\n{text[:3000]}"
        )
        result = call_claude_json(prompt)
        if isinstance(result, dict) and result.get("title"):
            return result

        # Last resort: return raw text so user can at least paste manually
        return {"title": "", "company": "", "description": text[:4000]}

    except Exception as exc:
        log.warning("Requests scrape failed: %s", exc)
        return {"title": "", "company": "", "description": "", "error": str(exc)}


@app.route("/api/scrape", methods=["POST"])
def api_scrape():
    data = request.get_json() or {}
    url = data.get("url", "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400

    parsed = urlparse(url)
    domain = parsed.netloc.lower()

    # Validate LinkedIn URLs — must point to a specific job posting, not the feed
    if "linkedin.com" in domain:
        path = parsed.path.lower()
        # Valid LinkedIn job URLs: /jobs/view/<id or slug-with-id>
        # Invalid: /jobs/, /jobs/search/, /feed/, /notifications/, etc.
        if not re.search(r"/jobs/view/", path):
            return jsonify({
                "title": "", "company": "", "description": "",
                "error": (
                    "This looks like the LinkedIn jobs feed, not a specific job posting. "
                    "Please open an individual job listing and copy its URL. "
                    "It should look like: linkedin.com/jobs/view/job-title-at-company-1234567890"
                ),
            })

    # JS-rendered portals: use Playwright only — requests would get the same empty shell
    js_only = any(d in domain for d in ["linkedin.com", "dice.com", "workday.com", "myworkdayjobs.com"])

    if js_only:
        result = _scrape_with_playwright(url, domain)
        # Don't fall back to requests for JS-only sites; it would return login-wall garbage
    else:
        # Server-rendered pages: requests + JSON-LD + Claude is fast and reliable
        result = _scrape_with_requests(url)
        # If requests came up empty (e.g. some hybrid JS page), try Playwright
        if not result.get("description"):
            pw_result = _scrape_with_playwright(url, domain)
            result = {
                "title":       pw_result.get("title")       or result.get("title", ""),
                "company":     pw_result.get("company")     or result.get("company", ""),
                "description": pw_result.get("description") or result.get("description", ""),
            }

    return jsonify(result)


@app.route("/api/upload-resume", methods=["POST"])
def api_upload_resume():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    f = request.files["file"]
    if not f.filename.lower().endswith(".docx"):
        return jsonify({"error": "Only .docx files are supported"}), 400
    save_path = UPLOAD_DIR / f.filename
    f.save(str(save_path))
    return jsonify({
        "key": "uploaded",
        "path": str(save_path),
        "name": f.filename,
        "label": "Uploaded Resume",
        "exists": True,
    })


@app.route("/api/analyze", methods=["POST"])
def api_analyze():
    data = request.get_json() or {}
    job_url = data.get("job_url", "")
    job_title = data.get("job_title", "")
    job_company = data.get("job_company", "")
    job_description = data.get("job_description", "")
    resume_key = data.get("resume_key", "qa")
    resume_path_override = data.get("resume_path")

    config = load_config()
    if resume_path_override:
        resume_path = resume_path_override
    else:
        resume_map = config.get("resume_map", {})
        resume_path = resume_map.get(resume_key, config.get("resume_path", ""))

    if not resume_path or not Path(resume_path).exists():
        return jsonify({"error": f"Resume not found: {resume_path}"}), 400

    resume_text = read_resume(resume_path)
    original_summary = extract_original_summary(resume_path)
    job1_orig, job2_orig, job1_tech_orig, job2_tech_orig = extract_original_bullets(resume_path)

    job = {
        "title": job_title,
        "company": job_company,
        "url": job_url,
        "full_description": job_description,
        "description_snippet": job_description[:500],
    }

    errors: list[str] = []

    # Single merged call: ATS analysis + tailoring combined (was 2 separate calls)
    combined: dict = {}
    try:
        combined = tailor_resume(resume_text, job)
    except Exception as e:
        errors.append(f"analyze+tailor: {e}")

    if errors:
        log.warning("Analysis warnings: %s", errors)

    match_pct_before: float = combined.get("match_percentage", 0)
    log.info("ATS+tailor: %.0f%% pre-tailor, %d missing keywords",
             match_pct_before, len(combined.get("missing_keywords", [])))

    # Split combined result: ATS fields go to ats_analysis, tailor fields go to tailored
    ats_fields = {"match_percentage", "matched_keywords", "missing_keywords", "top_skills", "recommendation", "ats_score"}
    ats_analysis = {k: v for k, v in combined.items() if k in ats_fields}
    tailored = {k: v for k, v in combined.items() if k not in ats_fields}

    return jsonify({
        "tailored": tailored,
        "ats_analysis": ats_analysis,
        "resume_path": resume_path,
        "resume_name": Path(resume_path).name,
        "original_summary": original_summary,
        "job1_original_bullets": job1_orig,
        "job2_original_bullets": job2_orig,
        "job1_tech_original_bullets": job1_tech_orig,
        "job2_tech_original_bullets": job2_tech_orig,
        "match_pct_before": match_pct_before,
    })


def _force_add_keywords_to_docx(output_path: Path, missing_keywords: list[str]) -> None:
    """Directly append missing keywords to the Skills section of the docx — no Claude needed."""
    from docx import Document
    from agent_resume import SKILLS_MARKERS, _para_is_section_heading, _set_paragraph_text

    if not missing_keywords:
        return

    doc = Document(str(output_path))
    in_skills = False
    for para in doc.paragraphs:
        if _para_is_section_heading(para, SKILLS_MARKERS):
            in_skills = True
            continue
        if in_skills and para.text.strip():
            existing = para.text.strip()
            existing_lower = existing.lower()
            new_kw = [k for k in missing_keywords if k.lower() not in existing_lower]
            if new_kw:
                sep = " | " if " | " in existing or "|" in existing else ", "
                _set_paragraph_text(para, existing + sep + sep.join(new_kw))
                log.info("Force-injected %d keywords into Skills section", len(new_kw))
            doc.save(str(output_path))
            return

    log.warning("Skills section not found — could not force-inject keywords")


def _ats_score(resume_text: str, jd: str) -> dict:
    """Claude-based post-generation ATS score."""
    prompt = (
        "Score this resume against the job description for ATS keyword coverage. "
        "Return JSON only (no markdown):\n"
        '{"match_percentage":<0-100>,"matched_keywords":["all found"],"missing_keywords":["all absent"]}\n\n'
        f"Job Description:\n{jd[:4000]}\n\nResume:\n{resume_text[:5000]}"
    )
    result = call_claude_json(prompt)
    return result if isinstance(result, dict) else {}


@app.route("/api/generate", methods=["POST"])
def api_generate():
    data = request.get_json() or {}
    resume_path = data.get("resume_path", "")
    tailored = data.get("tailored", {})
    job = data.get("job", {})

    if not resume_path or not Path(resume_path).exists():
        return jsonify({"error": f"Resume not found: {resume_path}"}), 400

    folder_raw = f"{job.get('title', 'Job')} - {job.get('company', 'Company')}"
    folder_safe = re.sub(r'[<>:"/\\|?*]', "_", folder_raw).strip()
    output_dir = Path("output/resumes") / folder_safe
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / Path(resume_path).name

    patch_docx(resume_path, tailored, job, output_path)

    # Score the generated document locally — no Claude call needed
    ats_score_after: dict = {}
    jd = data.get("job_description", "")
    try:
        if jd:
            final_text = read_resume(str(output_path))
            ats_score_after = _ats_score(final_text, jd)
            pct = ats_score_after.get("match_percentage", 0)
            log.info("Generated resume ATS score (local): %.0f%%", pct)

            # If below 75%, force-inject missing keywords and re-score
            if pct < 75 and ats_score_after.get("missing_keywords"):
                log.info("Score below 75%% — force-injecting %d keywords", len(ats_score_after["missing_keywords"]))
                _force_add_keywords_to_docx(output_path, ats_score_after["missing_keywords"])
                final_text = read_resume(str(output_path))
                ats_score_after = _ats_score(final_text, jd)
                log.info("After force-inject: %.0f%%", ats_score_after.get("match_percentage", 0))
    except Exception as exc:
        log.warning("Post-generation ATS score failed: %s", exc)

    return jsonify({
        "output_path": str(output_path),
        "filename": output_path.name,
        "folder": folder_safe,
        "ats_after": ats_score_after,
    })


@app.route("/api/download")
def api_download():
    file_path = request.args.get("path", "")
    if not file_path:
        return "No path specified", 400
    p = Path(file_path)
    if not p.exists():
        return "File not found", 404
    return send_file(str(p.resolve()), as_attachment=True, download_name=p.name)


@app.route("/api/apply", methods=["POST"])
def api_apply():
    data = request.get_json() or {}
    session_id = f"apply_{int(time.time() * 1000)}"
    q: Queue = Queue()
    _apply_sessions[session_id] = q

    config = load_config()
    applicant = get_applicant(config)

    t = threading.Thread(target=_run_apply_thread, args=(q, data, applicant), daemon=True)
    t.start()

    return jsonify({"session_id": session_id})


@app.route("/api/apply-status/<session_id>")
def api_apply_status(session_id: str):
    q = _apply_sessions.get(session_id)
    if not q:
        return jsonify({"error": "Session not found"}), 404

    def generate():
        while True:
            try:
                msg = q.get(timeout=120)
                yield f"data: {json.dumps(msg)}\n\n"
                if msg.get("done"):
                    _apply_sessions.pop(session_id, None)
                    break
            except Empty:
                yield f'data: {json.dumps({"status":"timeout","msg":"Timed out waiting for browser","done":True})}\n\n'
                _apply_sessions.pop(session_id, None)
                break

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


def _run_apply_thread(q: Queue, data: dict, applicant: dict) -> None:
    """Background thread: launches Playwright and streams progress via queue."""
    try:
        job_url = data.get("job_url", "")
        resume_path = data.get("resume_path", "")
        job_title = data.get("job_title", "")
        company = data.get("company", "")

        platform = detect_platform(job_url)
        q.put({"status": "starting", "msg": f"Platform detected: {platform.upper()}", "done": False})

        job_data = {
            "job_url": job_url,
            "job_title": job_title,
            "company": company,
            "resume_path": resume_path,
        }

        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            q.put({
                "status": "error",
                "msg": "Playwright not installed. Run: pip3 install playwright && playwright install chromium",
                "done": True,
            })
            return

        SESSION_FILE = Path("output/linkedin_session.json")
        q.put({"status": "browser", "msg": "Launching browser window...", "done": False})

        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=False, slow_mo=150)
            ctx_kwargs: dict = {
                "viewport": {"width": 1280, "height": 800},
                "user_agent": (
                    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
                ),
            }
            if SESSION_FILE.exists():
                ctx_kwargs["storage_state"] = str(SESSION_FILE)
                q.put({"status": "session", "msg": "LinkedIn session loaded", "done": False})
            else:
                q.put({"status": "warn", "msg": "No saved LinkedIn session — manual login may be required", "done": False})

            context = browser.new_context(**ctx_kwargs)
            page = context.new_page()

            try:
                from playwright_stealth import Stealth
                Stealth().apply_stealth_sync(page)
            except Exception:
                pass

            q.put({"status": "navigating", "msg": f"Opening: {job_url[:70]}...", "done": False})

            handler = PLATFORM_HANDLERS.get(platform, apply_generic)
            result = handler(page, job_data, applicant)
            page.close()

        q.put({
            "status": result.get("status", "unknown"),
            "msg": result.get("reason", ""),
            "platform": platform,
            "done": True,
        })

    except Exception as exc:
        log.exception("Apply thread error")
        q.put({"status": "error", "msg": str(exc)[:300], "done": True})


if __name__ == "__main__":
    log.info("Starting Job Hunter UI at http://localhost:5001")
    app.run(debug=False, port=5001, use_reloader=False, threaded=True)
